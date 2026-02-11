import os
import re
import pdfplumber
import docx
from pypdf import PdfReader

def clean_text(text):
    """
    Cleans raw text by removing excessive whitespace, 
    weird artifacts, and standardized spacing.
    """
    if not text:
        return ""
    
    # 1. Replace zero-width spaces and non-breaking spaces
    text = text.replace('\u200b', '').replace('\xa0', ' ')
    
    # 2. Remove multiple spaces and newlines
    text = re.sub(r'\s+', ' ', text)
    
    # 3. Remove weird control characters
    text = re.sub(r'[^\x00-\x7F]+', ' ', text)
    
    return text.strip()

def extract_from_pdf(filepath):
    """
    High-quality extraction using pdfplumber (best for layout) 
    with a fallback to pypdf (faster/standard).
    """
    text = ""
    try:
        # METHOD A: pdfplumber (Best for keeping words separate)
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
                    
    except Exception as e:
        print(f"Warning: pdfplumber failed ({e}), switching to pypdf fallback.")
        
        # METHOD B: pypdf (Fallback)
        try:
            reader = PdfReader(filepath)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except Exception as e2:
            print(f"Error: Both PDF extraction methods failed. {e2}")
            return ""

    return clean_text(text)

def extract_from_docx(filepath):
    """
    Extracts text from Word documents including tables.
    """
    try:
        doc = docx.Document(filepath)
        full_text = []
        
        # 1. Extract Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
                
        # 2. Extract Tables (Important for Strategy Docs)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
                    
        return clean_text(" ".join(full_text))
        
    except Exception as e:
        print(f"Error reading DOCX file: {e}")
        return ""

def extract_text(filepath):
    """
    Main entry point. Detects file type and routes to correct extractor.
    """
    if not os.path.exists(filepath):
        print(f"Error: File not found at {filepath}")
        return ""

    ext = os.path.splitext(filepath)[1].lower()
    
    if ext == ".pdf":
        return extract_from_pdf(filepath)
    elif ext in [".docx", ".doc"]:
        return extract_from_docx(filepath)
    elif ext == ".txt":
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                return clean_text(f.read())
        except Exception as e:
            print(f"Error reading text file: {e}")
            return ""
    else:
        print(f"Unsupported file format: {ext}")
        return ""

# Debugging / Testing (Only runs if you execute this file directly)
if __name__ == "__main__":
    test_path = "sample_data/strategy.pdf" # Replace with a real file to test
    if os.path.exists(test_path):
        print(extract_text(test_path)[:500]) # Print first 500 charsexit